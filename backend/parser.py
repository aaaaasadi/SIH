import io
import re
from typing import Tuple, Dict, Any, List
import pypdf
import docx

class ResumeParser:
    """
    Extracts text from PDF, DOCX, and TXT files, and segments text into structured resume components.
    """

    @staticmethod
    def extract_text(file_bytes: bytes, filename: str) -> str:
        name_lower = filename.lower()
        if name_lower.endswith(".pdf"):
            return ResumeParser._extract_pdf(file_bytes)
        elif name_lower.endswith(".docx"):
            return ResumeParser._extract_docx(file_bytes)
        elif name_lower.endswith(".txt"):
            return ResumeParser._extract_txt(file_bytes)
        else:
            raise ValueError(f"Unsupported file extension in '{filename}'. Only .pdf, .docx, and .txt are supported.")

    @staticmethod
    def _extract_pdf(file_bytes: bytes) -> str:
        try:
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            if len(reader.pages) == 0:
                raise ValueError("PDF file has 0 pages or is empty.")
            
            full_text = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    full_text.append(page_text)
            
            extracted = "\n".join(full_text).strip()
            if not extracted:
                raise ValueError("PDF contains no extractable text layer (it may be a scanned image).")
            
            return ResumeParser._clean_text(extracted)
        except Exception as e:
            if "scanned" in str(e).lower() or "0 pages" in str(e).lower():
                raise e
            raise ValueError(f"Failed to parse PDF document: {str(e)}")

    @staticmethod
    def _extract_docx(file_bytes: bytes) -> str:
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            full_text = []
            
            # Paragraphs
            for p in doc.paragraphs:
                if p.text.strip():
                    full_text.append(p.text.strip())
            
            # Tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))
                        
            extracted = "\n".join(full_text).strip()
            if not extracted:
                raise ValueError("DOCX document contains no text.")
            return ResumeParser._clean_text(extracted)
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX document: {str(e)}")

    @staticmethod
    def _extract_txt(file_bytes: bytes) -> str:
        for enc in ['utf-8', 'latin-1', 'cp1252', 'ascii']:
            try:
                text = file_bytes.decode(enc).strip()
                if text:
                    return ResumeParser._clean_text(text)
            except Exception:
                continue
        raise ValueError("Could not decode text file with standard encodings.")

    @staticmethod
    def _clean_text(text: str) -> str:
        # Normalize carriage returns and non-breaking spaces
        text = text.replace('\r\n', '\n').replace('\r', '\n').replace('\xa0', ' ')
        # Remove null bytes or weird unprintable control characters
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
        # Collapse excessive blank lines
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()

    @staticmethod
    def extract_contact_info(text: str) -> Dict[str, str]:
        contact = {
            "name": "Candidate",
            "email": "",
            "phone": "",
            "location": "",
            "linkedin": "",
            "github": ""
        }

        # 1. Email extraction
        email_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
        if email_match:
            contact["email"] = email_match.group(0)

        # 2. Phone extraction (handles (+91) 98765-43210, +91 98765 43210, (555) 123-4567, etc.)
        phone_match = re.search(r'(?:(?:\(\+\d{1,3}\)|\+\d{1,3})[-.\s]?)?(?:\(?\d{3,5}\)?[-.\s]?)?\d{3,5}[-.\s]?\d{4,5}', text)
        if phone_match:
            contact["phone"] = phone_match.group(0).strip()

        # 3. LinkedIn extraction
        linkedin_match = re.search(r'(?:https?://)?(?:www\.)?linkedin\.com/in/([a-zA-Z0-9_-]+)', text, re.IGNORECASE)
        if linkedin_match:
            contact["linkedin"] = f"linkedin.com/in/{linkedin_match.group(1)}"

        # 4. GitHub extraction
        github_match = re.search(r'(?:https?://)?(?:www\.)?github\.com/([a-zA-Z0-9_-]+)', text, re.IGNORECASE)
        if github_match:
            contact["github"] = f"github.com/{github_match.group(1)}"

        # 5. Location extraction (e.g. Bengaluru, India / San Francisco, CA / Remote)
        loc_match = re.search(r'([A-Z][a-zA-Z\s]+,\s*[A-Z]{2}(?:\s+\d{5})?|[A-Z][a-zA-Z\s]+,\s*(?:USA|Canada|India|UK|Germany|Remote))', text)
        if loc_match:
            contact["location"] = loc_match.group(1).strip()

        # 6. Name extraction (first 1-4 lines)
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        for line in lines[:5]:
            clean_line = re.sub(r'^(?:Name|Candidate Name|Full Name):\s*', '', line, flags=re.IGNORECASE).strip()
            if "@" in clean_line or "http" in clean_line or re.search(r'\d{3}', clean_line) or len(clean_line) > 50:
                continue
            words = clean_line.split()
            if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if w.isalpha()):
                contact["name"] = clean_line
                break

        if contact["name"] == "Candidate" and lines:
            first_line = re.sub(r'^(?:Name|Candidate Name|Full Name):\s*', '', lines[0].split('|')[0], flags=re.IGNORECASE).strip()
            if 2 <= len(first_line.split()) <= 4:
                contact["name"] = first_line

        return contact
