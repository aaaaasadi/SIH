import os
import io
import sys
import docx
import pypdf
import requests
import json

# Force UTF-8 on Windows
if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

BASE_URL = "http://127.0.0.1:8080"

def create_sample_docx() -> bytes:
    doc = docx.Document()
    doc.add_heading("Priya Sharma", 0)
    doc.add_paragraph("priya.sharma@email.com | (+91) 98765-43210 | Bengaluru, India | linkedin.com/in/priyasharma | github.com/priyasharma")
    
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph("Software Engineer with 5 years of experience in backend development, cloud infrastructure, and API design. Skilled in Python, Java, and AWS, with a proven record of improving system performance, reducing latency, and delivering scalable microservices. Strong background in Agile development, CI/CD pipelines, and cross-functional collaboration.")
    
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Programming Languages: Python, Java, JavaScript, SQL\nFrameworks and Libraries: Django, Spring Boot, React, Node.js\nCloud Platforms: AWS (EC2, S3, Lambda, RDS), GCP, Azure\nDatabases: MySQL, PostgreSQL, MongoDB, Redis\nDevOps and Tools: Docker, Kubernetes, Jenkins, Git, Terraform, CI/CD\nOther: REST API Design, Microservices Architecture, Agile/Scrum, Unit Testing, System Design")
    
    doc.add_heading("Professional Experience", level=1)
    doc.add_paragraph("Software Engineer II — Infosys Technologies (June 2021 - Present)")
    doc.add_paragraph("• Developed and maintained RESTful APIs using Python and Django, supporting over 50,000 daily active users.")
    doc.add_paragraph("• Reduced average API response time by 35% by optimizing database queries and implementing Redis caching.")
    doc.add_paragraph("• Led migration of monolithic application to microservices architecture on AWS, improving deployment frequency by 40%.")
    doc.add_paragraph("• Implemented automated CI/CD pipelines using Jenkins and Docker, reducing deployment time from 2 hours to 15 minutes.")
    doc.add_paragraph("• Collaborated with a cross-functional team of 8 engineers in an Agile Scrum environment to deliver features on a two-week sprint cycle.")
    doc.add_paragraph("• Mentored 3 junior engineers on best practices in code review, unit testing, and system design.")
    
    doc.add_paragraph("Software Engineer — Wipro Limited (July 2019 - May 2021)")
    doc.add_paragraph("• Built backend services in Java and Spring Boot for an e-commerce order management system processing 10,000 orders per day.")
    doc.add_paragraph("• Designed and implemented a MySQL database schema, improving query performance by 25%.")
    doc.add_paragraph("• Wrote unit and integration tests using JUnit, increasing code coverage from 60% to 90%.")
    doc.add_paragraph("• Participated in daily stand-ups, sprint planning, and retrospectives as part of an Agile development team.")
    
    doc.add_heading("Education", level=1)
    doc.add_paragraph("Bachelor of Technology in Computer Science and Engineering — Visvesvaraya Technological University, Belagavi, India (Graduated May 2019 | CGPA: 8.7/10.0)")

    doc.add_heading("Certifications", level=1)
    doc.add_paragraph("AWS Certified Solutions Architect - Associate (2022) | CKAD (2021) | Python Institute PCEP (2020)")

    doc.add_heading("Projects", level=1)
    doc.add_paragraph("Real-Time Chat Application: Node.js, Socket.io, MongoDB for 1,000 concurrent users.\nPersonal Finance Tracker: Django and React for expense tracking and budget analysis.")
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def create_sample_pdf() -> bytes:
    # Build a valid standard PDF with text stream for Priya Sharma
    pdf_content = (
        b"%PDF-1.4\n"
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>\nendobj\n"
        b"4 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n"
        b"5 0 obj\n<< /Length 480 >>\nstream\n"
        b"BT\n/F1 14 Tf\n50 720 Td\n(Priya Sharma) Tj\n"
        b"/F1 10 Tf\n0 -20 Td\n(priya.sharma@email.com | +91-98765-43210 | Bengaluru, India | linkedin.com/in/priyasharma) Tj\n"
        b"0 -30 Td\n(Professional Summary) Tj\n"
        b"0 -15 Td\n(Software Engineer with 5 years experience in Python, Java, AWS microservices, and Redis.) Tj\n"
        b"0 -30 Td\n(Professional Experience) Tj\n"
        b"0 -15 Td\n(Software Engineer II - Infosys Technologies) Tj\n"
        b"0 -15 Td\n(Developed RESTful APIs in Python/Django for 50,000 daily active users, optimizing latency by 35%.) Tj\n"
        b"0 -15 Td\n(Implemented Docker CI/CD pipelines reducing deployment time by 40%.) Tj\n"
        b"0 -30 Td\n(Technical Skills) Tj\n"
        b"0 -15 Td\n(Python, Java, Django, Spring Boot, AWS, Docker, Kubernetes, MySQL, Redis, Git) Tj\n"
        b"ET\nendstream\nendobj\nxref\n0 6\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n0000000115 00000 n \n0000000244 00000 n \n0000000318 00000 n \ntrailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n850\n%%EOF\n"
    )
    return pdf_content

def test_api():
    print("--- 1. Testing Health Endpoint ---")
    r = requests.get(f"{BASE_URL}/api/health")
    print(f"Health Response ({r.status_code}):", r.json())
    assert r.status_code == 200

    print("\n--- 2. Testing DOCX Resume Upload with Job Description ---")
    docx_bytes = create_sample_docx()
    files = {
        "file": ("priya_sharma_resume.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    }
    data = {
        "job_description": "We are seeking a Software Engineer with expertise in Python, Java, Django, Spring Boot, AWS, Docker, Kubernetes, MySQL, Redis, and CI/CD."
    }
    r = requests.post(f"{BASE_URL}/api/resume/analyze", files=files, data=data)
    print(f"Analyze DOCX Response ({r.status_code}):")
    assert r.status_code == 200
    res = r.json()
    print(f"  • Resume Score: {res['resume_score']}/100")
    print(f"  • ATS Score: {res['ats_score']}/100")
    print(f"  • Keyword Alignment: {res['keyword_alignment']}%")
    print(f"  • Interview Readiness: {res['interview_readiness']}/100")
    print(f"  • Candidate: {res['parsed_resume']['candidate']['name']} ({res['parsed_resume']['candidate']['email']})")
    print(f"  • Matching Keywords: {res['matching_keywords']}")
    print(f"  • Missing Keywords: {res['missing_keywords']}")
    print(f"  • Strengths: {res['strengths']}")
    print(f"  • ATS Issues: {res['ats_issues']}")
    assert res['ats_score'] >= 85
    assert res['parsed_resume']['candidate']['name'] == "Priya Sharma"

    print("\n--- 3. Testing PDF (.pdf) Resume Upload ---")
    pdf_bytes = create_sample_pdf()
    files = {
        "file": ("priya_sharma.pdf", pdf_bytes, "application/pdf")
    }
    data = {
        "job_description": "Software Engineer with Python, Java, AWS, Docker, and Redis."
    }
    r = requests.post(f"{BASE_URL}/api/resume/analyze", files=files, data=data)
    print(f"Analyze PDF Response ({r.status_code}):")
    assert r.status_code == 200
    res_pdf = r.json()
    print(f"  • Resume Score: {res_pdf['resume_score']}/100")
    print(f"  • ATS Score: {res_pdf['ats_score']}/100")
    print(f"  • Candidate: {res_pdf['parsed_resume']['candidate']['name']} ({res_pdf['parsed_resume']['candidate']['email']})")
    assert res_pdf['parsed_resume']['candidate']['name'] == "Priya Sharma"

    print("\n--- 4. Testing Plain Text (.txt) Resume Upload ---")
    txt_content = """
    Aarav Sharma
    aarav.sharma@sjsu.edu | (408) 555-0144 | San Jose, CA | linkedin.com/in/aaravsharma
    
    Education
    B.S. in Computer Science, San Jose State University, GPA: 3.82 (2026)
    
    Technical Skills
    Java, Python, SQL, React, Node.js, Git, Docker, REST API
    
    Projects
    • Smart Campus Resolution System: Built full-stack issue tracking system in React and FastAPI.
    • Developed automated unit tests with Pytest.
    """
    files = {
        "file": ("aarav_resume.txt", txt_content.encode("utf-8"), "text/plain")
    }
    data = {
        "job_description": "Software Engineer role requiring Java, Python, SQL, REST API, Docker, AWS, Kubernetes."
    }
    r = requests.post(f"{BASE_URL}/api/resume/analyze", files=files, data=data)
    print(f"Analyze TXT Response ({r.status_code}):")
    assert r.status_code == 200
    res_txt = r.json()
    print(f"  • Resume Score: {res_txt['resume_score']}/100")
    print(f"  • ATS Score: {res_txt['ats_score']}/100")
    print(f"  • Matching Keywords: {res_txt['matching_keywords']}")
    print(f"  • Missing Keywords: {res_txt['missing_keywords']}")

    print("\n--- 5. Testing Recent Analyses DB Retrieval ---")
    r = requests.get(f"{BASE_URL}/api/analyses/recent")
    assert r.status_code == 200
    analyses = r.json()
    print(f"Recent Analyses ({r.status_code}):", len(analyses), "records found in SQLite DB.")

    print("\n--- 6. Testing Static Web Frontend Serving ---")
    r = requests.get(f"{BASE_URL}/")
    assert r.status_code == 200
    print(f"Index.html ({r.status_code}): Size {len(r.text)} chars")
    assert "<title>CareerAI" in r.text

    print("\n[SUCCESS] ALL INTEGRATION TESTS PASSED SUCCESSFULLY!")

if __name__ == "__main__":
    test_api()
